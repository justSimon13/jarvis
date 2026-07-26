"""
Erzeugt PDF- oder Word-Dokumente aus jeder in SeiteView.vue anzeigbaren Seite —
Wurzel-Seite eines Projekts/Todos/Kontakts oder eine einzelne Unterseite (lokales
local_data-Schema, siehe local_data.py: projekte/todos/kontakte/seiten-Tabellen,
_sections_for() nutzt local_data.get_seite_view() — dieselbe Funktion wie das
Frontend zum Anzeigen).

Ein gemeinsamer, bewusst einfacher Markdown-Block-Parser (_parse_blocks) speist
zwei Renderer — reportlab für PDF, python-docx für Word. Kein HTML-Zwischenschritt
und keine Systemabhängigkeiten (anders als z.B. weasyprint/pandoc): beide Libraries
sind pip-only, passt zu "kein Docker"/einfacher Deploy. Deckt Überschriften (#/##/###),
Absätze, Bullet-Listen, horizontale Linien (---), Zitate (>), GFM-Pipe-Tabellen und
Inline-Fett/Kursiv ab — kein CommonMark-kompletter Parser, aber reicht für JARVIS'
Seiten-Inhalte (Notizen/PRDs/Recherche). Nicht abgedeckt: nummerierte Listen,
verschachtelte Listen, Code-Blöcke, Links, Bilder.
"""
import base64
import io
import re

import local_data

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable, Table, TableStyle,
)

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)")
_HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
_QUOTE_RE = re.compile(r"^>\s?(.*)")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
_TABLE_SEP_RE = re.compile(r"^\|[\s:|-]+\|$")
# Reihenfolge in der Alternation ist entscheidend: **fett** muss vor *kursiv*
# geprüft werden, sonst würde das Bold-Sternchenpaar als zwei Kursiv-Marker
# fehlinterpretiert (Regex-Alternation probiert Alternativen an jeder Position
# links-nach-rechts, erste passende gewinnt).
_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_")
_SAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9_\-]+")


def _split_table_row(line: str) -> list[str]:
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [cell.strip() for cell in inner.split("|")]


def _parse_blocks(markdown_text: str) -> list[dict]:
    """Zerlegt Markdown in Blöcke: Überschrift ({"type":"h","level":N,"text":...}),
    Bullet-Liste ({"type":"ul","items":[...]}), horizontale Linie ({"type":"hr"}),
    Zitat ({"type":"quote","text":...}), Tabelle ({"type":"table","rows":[[...],...]},
    erste Zeile = Kopf), sonst Absatz ({"type":"p","text":...})."""
    blocks = []
    lines = (markdown_text or "").splitlines()
    paragraph_buf = []

    def flush_paragraph():
        text = " ".join(paragraph_buf).strip()
        if text:
            blocks.append({"type": "p", "text": text})
        paragraph_buf.clear()

    i = 0
    n = len(lines)
    while i < n:
        stripped = lines[i].strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph()
            blocks.append({"type": "h", "level": len(heading_match.group(1)), "text": heading_match.group(2).strip()})
            i += 1
            continue

        if _HR_RE.match(stripped):
            flush_paragraph()
            blocks.append({"type": "hr"})
            i += 1
            continue

        if _TABLE_ROW_RE.match(stripped) and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1].strip()):
            flush_paragraph()
            rows = [_split_table_row(stripped)]
            i += 2  # Kopfzeile + Trennzeile (|---|---|) überspringen
            while i < n and _TABLE_ROW_RE.match(lines[i].strip()):
                rows.append(_split_table_row(lines[i].strip()))
                i += 1
            blocks.append({"type": "table", "rows": rows})
            continue

        quote_match = _QUOTE_RE.match(stripped)
        if quote_match:
            flush_paragraph()
            quote_lines = [quote_match.group(1).strip()]
            while i + 1 < n and _QUOTE_RE.match(lines[i + 1].strip()):
                i += 1
                quote_lines.append(_QUOTE_RE.match(lines[i].strip()).group(1).strip())
            blocks.append({"type": "quote", "text": " ".join(quote_lines).strip()})
            i += 1
            continue

        bullet_match = _BULLET_RE.match(stripped)
        if bullet_match:
            flush_paragraph()
            items = [bullet_match.group(1).strip()]
            while i + 1 < n and _BULLET_RE.match(lines[i + 1].strip()):
                i += 1
                items.append(_BULLET_RE.match(lines[i].strip()).group(1).strip())
            blocks.append({"type": "ul", "items": items})
            i += 1
            continue

        paragraph_buf.append(stripped)
        i += 1

    flush_paragraph()
    return blocks


def _split_inline(text: str) -> list[tuple[str, bool, bool]]:
    """Zerlegt Text an **fett**/*kursiv*/_kursiv_-Markierungen in (Text, ist_fett,
    ist_kursiv)-Segmente. Kein verschachteltes/kombiniertes ***fett+kursiv*** —
    reicht für JARVIS' Notizen-Content."""
    segments = []
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False, False))
        if m.group(1) is not None:
            segments.append((m.group(1), True, False))
        else:
            segments.append((m.group(2) if m.group(2) is not None else m.group(3), False, True))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False, False))
    return segments or [(text, False, False)]


def _blocks_for_section(title: str, markdown_text: str) -> list[dict]:
    """_parse_blocks(), aber verwirft eine führende Überschrift, die nur den
    Section-Titel wiederholt — Seiten fangen ihren Inhalt oft mit '# <eigener
    Titel>' an, sonst stünde der Titel doppelt im Dokument (einmal als
    Section-Heading, einmal als erster Content-Block)."""
    blocks = _parse_blocks(markdown_text)
    if blocks and blocks[0]["type"] == "h" and blocks[0]["text"].strip().lower() == (title or "").strip().lower():
        blocks = blocks[1:]
    return blocks


def _collect_seite_tree(seite_id: int, depth: int = 1) -> list[tuple[str, str, int]]:
    """(titel, inhalt, tiefe) für eine Seite + alle Unterseiten, rekursiv,
    Tiefe-zuerst-Reihenfolge (passt zur Anzeigereihenfolge im Frontend)."""
    seite = local_data.get_seite(seite_id)
    if not seite:
        return []
    result = [(seite["titel"], seite.get("inhalt") or "", depth)]
    for kind in local_data.list_unterseiten(seite_id):
        result.extend(_collect_seite_tree(kind["id"], depth + 1))
    return result


def _sections_for(quelle_typ: str, quelle_id: int) -> tuple[str, str, list[tuple[str, str, int]]]:
    """Gibt (Dokumenttitel, Intro-Markdown, [(Abschnittstitel, Inhalt, Tiefe), ...]) zurück.
    quelle_typ: dieselben Werte wie überall sonst im System — 'projekte'/'todos'/'kontakte'
    (Wurzel-Seite eines Eintrags) oder 'seite' (eine einzelne Unterseite). Ein einziger
    Codepfad für alle vier über local_data.get_seite_view() — das ist exakt dieselbe
    Funktion, die auch SeiteView.vue im Frontend zum Anzeigen nutzt, also automatisch
    für jede dort sichtbare Seite nutzbar, nicht nur für Projekte."""
    view = local_data.get_seite_view(quelle_typ, quelle_id)
    if not view:
        raise ValueError(f"{quelle_typ} {quelle_id} nicht gefunden.")

    intro = view["inhalt"] or ""
    beschreibung = (view.get("meta") or {}).get("beschreibung")
    if beschreibung:
        intro = (beschreibung + "\n\n" + intro).strip()

    children = []
    for kind in view["unterseiten"]:
        children.extend(_collect_seite_tree(kind["id"]))

    return view["titel"] or "Dokument", intro, children


# ── PDF (reportlab) ───────────────────────────────────────────────────────────

def _escape_reportlab(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _markup_reportlab(text: str) -> str:
    parts = []
    for chunk, bold, italic in _split_inline(text):
        chunk = _escape_reportlab(chunk)
        if bold:
            chunk = f"<b>{chunk}</b>"
        if italic:
            chunk = f"<i>{chunk}</i>"
        parts.append(chunk)
    return "".join(parts)


def _render_pdf(doc_title: str, intro: str, children: list[tuple[str, str, int]]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=doc_title)
    styles = getSampleStyleSheet()
    heading_styles = {1: styles["Heading1"], 2: styles["Heading2"], 3: styles["Heading3"]}
    quote_style = ParagraphStyle("Quote", parent=styles["BodyText"], leftIndent=20, textColor=colors.HexColor("#555555"))

    def blocks_to_flowables(blocks):
        flowables = []
        for block in blocks:
            if block["type"] == "h":
                flowables.append(Paragraph(_escape_reportlab(block["text"]), heading_styles.get(block["level"], styles["Heading3"])))
            elif block["type"] == "p":
                flowables.append(Paragraph(_markup_reportlab(block["text"]), styles["BodyText"]))
            elif block["type"] == "quote":
                flowables.append(Paragraph(f"<i>{_markup_reportlab(block['text'])}</i>", quote_style))
            elif block["type"] == "hr":
                flowables.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey, spaceBefore=4, spaceAfter=4))
            elif block["type"] == "ul":
                flowables.append(ListFlowable(
                    [ListItem(Paragraph(_markup_reportlab(item), styles["BodyText"])) for item in block["items"]],
                    bulletType="bullet",
                ))
            elif block["type"] == "table":
                table_data = [[Paragraph(_markup_reportlab(cell), styles["BodyText"]) for cell in row] for row in block["rows"]]
                tbl = Table(table_data, hAlign="LEFT")
                tbl.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                flowables.append(tbl)
            flowables.append(Spacer(1, 6))
        return flowables

    story = [Paragraph(_escape_reportlab(doc_title), styles["Title"]), Spacer(1, 16)]
    story.extend(blocks_to_flowables(_blocks_for_section(doc_title, intro)))
    for title, content, depth in children:
        story.append(Spacer(1, 8))
        story.append(Paragraph(_escape_reportlab(title), heading_styles.get(min(depth, 3), styles["Heading3"])))
        story.extend(blocks_to_flowables(_blocks_for_section(title, content)))

    doc.build(story)
    return buf.getvalue()


# ── Word (python-docx) ────────────────────────────────────────────────────────

def _add_runs(paragraph, text: str) -> None:
    for chunk, bold, italic in _split_inline(text):
        run = paragraph.add_run(chunk)
        run.bold = bold
        run.italic = italic


def _add_horizontal_rule(document) -> None:
    """python-docx hat kein eingebautes HR-Element — Standard-Workaround: ein
    leerer Absatz mit unterem Rahmen (bottom border) via rohem OOXML."""
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _render_docx(doc_title: str, intro: str, children: list[tuple[str, str, int]]) -> bytes:
    document = Document()
    document.add_heading(doc_title, level=0)

    def add_blocks(blocks):
        for block in blocks:
            if block["type"] == "h":
                document.add_heading(block["text"], level=min(block["level"] + 1, 4))
            elif block["type"] == "p":
                _add_runs(document.add_paragraph(), block["text"])
            elif block["type"] == "quote":
                try:
                    paragraph = document.add_paragraph(style="Quote")
                except KeyError:
                    # Fallback falls das Default-Template die eingebaute "Quote"-Formatvorlage
                    # nicht enthält: manueller Einzug statt Formatvorlage.
                    paragraph = document.add_paragraph()
                    paragraph.paragraph_format.left_indent = Inches(0.3)
                _add_runs(paragraph, block["text"])
            elif block["type"] == "hr":
                _add_horizontal_rule(document)
            elif block["type"] == "ul":
                for item in block["items"]:
                    _add_runs(document.add_paragraph(style="List Bullet"), item)
            elif block["type"] == "table":
                rows = block["rows"]
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                try:
                    table.style = "Table Grid"
                except KeyError:
                    pass
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        cell_paragraph = table.cell(r, c).paragraphs[0]
                        _add_runs(cell_paragraph, cell_text)
                        if r == 0:
                            for run in cell_paragraph.runs:
                                run.bold = True

    add_blocks(_blocks_for_section(doc_title, intro))
    for title, content, depth in children:
        document.add_heading(title, level=min(depth + 1, 4))
        add_blocks(_blocks_for_section(title, content))

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def generate(quelle_typ: str, quelle_id: int, format: str) -> tuple[str, str, str]:
    """Gibt (filename, mime_type, data_base64) zurück. Wirft ValueError bei
    unbekannter Quelle/unbekanntem Format."""
    if format not in ("pdf", "docx"):
        raise ValueError(f"Unbekanntes Format: {format} (erlaubt: pdf, docx)")

    doc_title, intro, children = _sections_for(quelle_typ, quelle_id)

    if format == "pdf":
        data = _render_pdf(doc_title, intro, children)
        mime = "application/pdf"
    else:
        data = _render_docx(doc_title, intro, children)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    safe_name = _SAFE_FILENAME_RE.sub("_", doc_title).strip("_") or "dokument"
    filename = f"{safe_name}.{format}"
    return filename, mime, base64.b64encode(data).decode("ascii")
